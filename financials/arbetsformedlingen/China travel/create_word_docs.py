"""
Script to create Word documents for China trip
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

output_dir = r"c:\Users\samij\Desktop\Visma revesor info for old partner and new samisjackets and auronis vip\financials\arbetsformedlingen\China travel\Word Documents"
os.makedirs(output_dir, exist_ok=True)

def set_margins(doc, margin_cm=2):
    """Set document margins"""
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)

def add_signature_line(doc, name, personnummer=None):
    """Add a signature line"""
    p = doc.add_paragraph()
    p.add_run("\n" + "_" * 50)
    p = doc.add_paragraph()
    p.add_run(f"{name}")
    if personnummer:
        p.add_run(f" ({personnummer})")
    p = doc.add_paragraph()
    p.add_run("Datum: _______________")

# =============================================================================
# 1. FULLMAKT - Power of Attorney for Children
# =============================================================================
def create_fullmakt():
    doc = Document()
    set_margins(doc, 2)
    
    # Title
    title = doc.add_heading('FULLMAKT FÖR TILLFÄLLIG VÅRDNAD AV BARN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Under föräldrarnas utlandsresa 27 november – 12 december 2025')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Parents section
    doc.add_heading('FULLMAKTSGIVARE (FÖRÄLDRAR)', level=1)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    cells = table.rows[0].cells
    cells[0].text = "Far"
    cells[1].text = "Mor"
    
    cells = table.rows[1].cells
    cells[0].text = "Mohammad Sami Alsharef\n860217-5070\nTel: 0720147440 / 0735116961\ninfo@samisjackets.com"
    cells[1].text = "Jomana Alnablsi\n880102-5084\nTel: 0735165708\njomana.alnablsi@hotmail.com"
    
    cells = table.rows[2].cells
    cells[0].text = "Adress: Esbjergsgatan 1 Lgh 1302, 632 30 Eskilstuna"
    cells[0].merge(cells[1])
    
    doc.add_paragraph()
    
    # Caregivers
    doc.add_heading('FULLMAKTSTAGARE (TILLFÄLLIGA VÅRDNADSHAVARE)', level=1)
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    cells = table.rows[0].cells
    cells[0].text = "Emelie Nilsson\n19990408-5306\nTel: 0768997423"
    cells[1].text = "Mari Johansson\n19720522-1646\nTel: 0708900945"
    
    cells = table.rows[1].cells
    cells[0].text = "Storgatan 13, 64530 Strängnäs"
    cells[1].text = "Forsbomsgatan 18, Eskilstuna"
    
    doc.add_paragraph()
    
    # Children
    doc.add_heading('BARN SOM OMFATTAS', level=1)
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Namn"
    headers[1].text = "Personnummer"
    headers[2].text = "Ålder"
    headers[3].text = "Mobil"
    
    children = [
        ("Lilas Al Sharef", "160204-4040", "9 år", "0793392306"),
        ("Hoda Alshrif", "100217-7309", "15 år", "0735299675"),
        ("Alyaah Alshrif", "080925-9666", "17 år", "0735147846"),
    ]
    
    for i, child in enumerate(children):
        cells = table.rows[i+1].cells
        for j, value in enumerate(child):
            cells[j].text = value
    
    doc.add_paragraph()
    
    # Authorization
    doc.add_heading('FULLMAKTEN GER RÄTT ATT', level=1)
    
    rights = [
        "Söka sjukvård och ge samtycke till nödvändig medicinsk behandling",
        "Hämta ut receptbelagd medicin på apotek",
        "Kommunicera med skola och signera skolblanketter",
        "Anmäla frånvaro vid sjukdom",
        "Fatta vardagliga beslut för barnens bästa",
        "Ansvara för barnens dagliga omsorg, mat och boende"
    ]
    
    for right in rights:
        doc.add_paragraph(f"• {right}")
    
    doc.add_paragraph()
    
    # Period
    p = doc.add_paragraph()
    p.add_run("GILTIGHETSTID: ").bold = True
    p.add_run("27 november – 12 december 2025 (16 dagar)")
    
    p = doc.add_paragraph()
    p.add_run("ANLEDNING: ").bold = True
    p.add_run("Affärsresa till Kina")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures
    doc.add_heading('UNDERSKRIFTER', level=1)
    
    doc.add_paragraph("FÖRÄLDRAR:")
    
    table = doc.add_table(rows=3, cols=2)
    cells = table.rows[0].cells
    cells[0].text = "\n\n_________________________________"
    cells[1].text = "\n\n_________________________________"
    
    cells = table.rows[1].cells
    cells[0].text = "Mohammad Sami Alsharef (860217-5070)"
    cells[1].text = "Jomana Alnablsi (880102-5084)"
    
    cells = table.rows[2].cells
    cells[0].text = "Datum: _______________"
    cells[1].text = "Datum: _______________"
    
    doc.add_paragraph()
    doc.add_paragraph("TILLFÄLLIGA VÅRDNADSHAVARE:")
    
    table = doc.add_table(rows=3, cols=2)
    cells = table.rows[0].cells
    cells[0].text = "\n\n_________________________________"
    cells[1].text = "\n\n_________________________________"
    
    cells = table.rows[1].cells
    cells[0].text = "Emelie Nilsson (19990408-5306)"
    cells[1].text = "Mari Johansson (19720522-1646)"
    
    cells = table.rows[2].cells
    cells[0].text = "Datum: _______________"
    cells[1].text = "Datum: _______________"
    
    doc.save(os.path.join(output_dir, "FULLMAKT_Barn_27nov-12dec_2025.docx"))
    print("✅ FULLMAKT skapad!")

# =============================================================================
# 2. VÄNSKAPSAVTAL - Friendship Agreement
# =============================================================================
def create_vanskapsavtal():
    doc = Document()
    set_margins(doc, 2)
    
    title = doc.add_heading('VÄNSKAPSAVTAL OM BARNOMSORG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('27 november – 12 december 2025')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Introduction
    intro = doc.add_paragraph()
    intro.add_run("Detta avtal bekräftar att ").italic = False
    intro.add_run("Emelie Nilsson").bold = True
    intro.add_run(" och ")
    intro.add_run("Mari Johansson").bold = True
    intro.add_run(" frivilligt och av vänskap hjälper till med barnomsorg under föräldrarnas affärsresa till Kina. Detta är en vänskapstjänst utan ekonomisk ersättning.")
    
    doc.add_paragraph()
    
    # Parties
    doc.add_heading('PARTER', level=1)
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    cells = table.rows[0].cells
    cells[0].text = "FÖRÄLDRAR"
    cells[1].text = "VÄNNER SOM HJÄLPER"
    
    cells = table.rows[1].cells
    cells[0].text = "Mohammad Sami Alsharef (860217-5070)\nTel: 0720147440 / 0735116961\n\nJomana Alnablsi (880102-5084)\nTel: 0735165708"
    cells[1].text = "Emelie Nilsson (19990408-5306)\nTel: 0768997423\n\nMari Johansson (19720522-1646)\nTel: 0708900945"
    
    doc.add_paragraph()
    
    # Children
    doc.add_heading('BARNEN', level=1)
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Namn"
    headers[1].text = "Ålder"
    headers[2].text = "Mobil"
    
    children = [
        ("Lilas Al Sharef", "9 år", "0793392306"),
        ("Hoda Alshrif", "15 år", "0735299675"),
        ("Alyaah Alshrif", "17 år", "0735147846"),
    ]
    
    for i, child in enumerate(children):
        cells = table.rows[i+1].cells
        for j, value in enumerate(child):
            cells[j].text = value
    
    doc.add_paragraph()
    
    # Responsibilities
    doc.add_heading('EMELIE OCH MARI HJÄLPER TILL MED', level=1)
    
    tasks = [
        "Se till att barnen mår bra och är trygga",
        "Hjälpa Lilas till och från skolan",
        "Finnas tillgängliga för frågor och stöd",
        "Kontakta föräldrarna vid behov",
        "Vid sjukdom: kontakta 1177 eller sjukvård"
    ]
    
    for task in tasks:
        doc.add_paragraph(f"✓ {task}")
    
    doc.add_paragraph()
    
    # Thanks
    doc.add_heading('TACK', level=1)
    thank = doc.add_paragraph()
    thank.add_run("Vi är oerhört tacksamma för att ni ställer upp för vår familj. Att veta att våra barn är i trygga händer hos goda vänner betyder allt för oss. Tack från hjärtat!").italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures
    doc.add_heading('UNDERSKRIFTER', level=1)
    
    table = doc.add_table(rows=3, cols=2)
    cells = table.rows[0].cells
    cells[0].text = "\n_________________________________\nMohammad Sami Alsharef"
    cells[1].text = "\n_________________________________\nJomana Alnablsi"
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=3, cols=2)
    cells = table.rows[0].cells
    cells[0].text = "\n_________________________________\nEmelie Nilsson"
    cells[1].text = "\n_________________________________\nMari Johansson"
    
    p = doc.add_paragraph()
    p.add_run("\nDatum: _______________     Ort: Eskilstuna")
    
    doc.save(os.path.join(output_dir, "VANSKAPSAVTAL_Barnomsorg_27nov-12dec_2025.docx"))
    print("✅ VÄNSKAPSAVTAL skapad!")

# =============================================================================
# 3. NÖDKONTAKTER - Emergency Contacts (1 page)
# =============================================================================
def create_nodkontakter():
    doc = Document()
    set_margins(doc, 1.5)
    
    title = doc.add_heading('🆘 NÖDKONTAKTER - KINARESAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('27 nov – 12 dec 2025')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Emergency numbers
    doc.add_heading('NÖDNUMMER', level=1)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    data = [
        ("SOS Alarm", "112"),
        ("Sjukvårdsrådgivning", "1177"),
        ("Polis (ej akut)", "114 14"),
        ("Giftinformation", "010-456 67 00"),
    ]
    for i, (service, number) in enumerate(data):
        cells = table.rows[i].cells
        cells[0].text = service
        cells[1].text = number
        for cell in cells:
            cell.paragraphs[0].runs[0].bold = True if i == 0 else False
    
    doc.add_paragraph()
    
    # Parents
    doc.add_heading('FÖRÄLDRAR I KINA', level=1)
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    cells = table.rows[0].cells
    cells[0].text = "Mohammad Sami (Pappa)"
    cells[1].text = "0720147440 / 0735116961\ninfo@samisjackets.com"
    
    cells = table.rows[1].cells
    cells[0].text = "Jomana (Mamma)"
    cells[1].text = "0735165708\njomana.alnablsi@hotmail.com"
    
    doc.add_paragraph()
    
    # Children
    doc.add_heading('BARNEN', level=1)
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Barn"
    headers[1].text = "Personnummer"
    headers[2].text = "Mobil"
    
    children = [
        ("Lilas (9 år)", "160204-4040", "0793392306"),
        ("Hoda (15 år)", "100217-7309", "0735299675"),
        ("Alyaah (17 år)", "080925-9666", "0735147846"),
    ]
    
    for i, child in enumerate(children):
        cells = table.rows[i+1].cells
        for j, value in enumerate(child):
            cells[j].text = value
    
    doc.add_paragraph()
    
    # Address
    doc.add_heading('ADRESS', level=1)
    p = doc.add_paragraph()
    p.add_run("Esbjergsgatan 1 Lgh 1302, 632 30 Eskilstuna").bold = True
    
    doc.add_paragraph()
    
    # Hospital
    doc.add_heading('SJUKHUS', level=1)
    p = doc.add_paragraph()
    p.add_run("Mälarsjukhuset Eskilstuna").bold = True
    doc.add_paragraph("Kungsvägen, 631 88 Eskilstuna")
    doc.add_paragraph("Tel: 016-10 30 00")
    
    doc.save(os.path.join(output_dir, "NODKONTAKTER_Kinaresan.docx"))
    print("✅ NÖDKONTAKTER skapad!")

# =============================================================================
# 4. WEEKLY SCHEDULES - One page per week
# =============================================================================
def create_weekly_schedule(week_num, start_date, end_date, days_data):
    doc = Document()
    set_margins(doc, 1.5)
    
    title = doc.add_heading(f'ARBETSSCHEMA VECKA {week_num}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(f'{start_date} – {end_date}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Company info
    p = doc.add_paragraph()
    p.add_run("Sami's Jackets AB").bold = True
    p.add_run(" | Org.nr: 559489-5301 | Kungsgatan 24, 632 18 Eskilstuna")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Schedule table
    table = doc.add_table(rows=len(days_data)+1, cols=5)
    table.style = 'Table Grid'
    
    # Headers
    headers = table.rows[0].cells
    headers[0].text = "Dag"
    headers[1].text = "Datum"
    headers[2].text = "Emelie\n11:00-14:00"
    headers[3].text = "Mari\n14:00-17:00"
    headers[4].text = "Signatur"
    
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
    
    # Days
    for i, (day, date, emelie, mari) in enumerate(days_data):
        cells = table.rows[i+1].cells
        cells[0].text = day
        cells[1].text = date
        cells[2].text = emelie
        cells[3].text = mari
        cells[4].text = ""
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Summary
    doc.add_heading('VECKOSAMMANFATTNING', level=1)
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    cells = table.rows[0].cells
    cells[0].text = ""
    cells[1].text = "Emelie Nilsson"
    cells[2].text = "Mari Johansson"
    
    cells = table.rows[1].cells
    cells[0].text = "Timmar denna vecka"
    cells[1].text = "_____ tim"
    cells[2].text = "_____ tim"
    
    cells = table.rows[2].cells
    cells[0].text = "Timlön"
    cells[1].text = "110 SEK/tim"
    cells[2].text = "110 SEK/tim"
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures
    doc.add_heading('SIGNATUR', level=1)
    
    table = doc.add_table(rows=2, cols=3)
    
    cells = table.rows[0].cells
    cells[0].text = "\n\n_______________________\nEmelie Nilsson"
    cells[1].text = "\n\n_______________________\nMari Johansson"
    cells[2].text = "\n\n_______________________\nArbetsgivare"
    
    cells = table.rows[1].cells
    cells[0].text = "Datum: ___________"
    cells[1].text = "Datum: ___________"
    cells[2].text = "Datum: ___________"
    
    return doc

def create_all_weekly_schedules():
    # Week 48: 28 nov - 1 dec (Thu-Sun, but only workdays Thu-Fri)
    week48_days = [
        ("Torsdag", "28 nov", "3 tim", "3 tim"),
        ("Fredag", "29 nov", "3 tim", "3 tim"),
        ("Lördag", "30 nov", "Ledig", "Ledig"),
        ("Söndag", "1 dec", "Ledig", "Ledig"),
    ]
    doc = create_weekly_schedule(48, "28 nov", "1 dec 2025", week48_days)
    doc.save(os.path.join(output_dir, "SCHEMA_Vecka48_28nov-1dec.docx"))
    print("✅ SCHEMA Vecka 48 skapad!")
    
    # Week 49: 2-7 dec (Mon-Sun)
    week49_days = [
        ("Måndag", "2 dec", "3 tim", "3 tim"),
        ("Tisdag", "3 dec", "3 tim", "3 tim"),
        ("Onsdag", "4 dec", "3 tim", "3 tim"),
        ("Torsdag", "5 dec", "3 tim", "3 tim"),
        ("Fredag", "6 dec", "3 tim", "3 tim"),
        ("Lördag", "7 dec", "Ledig", "Ledig"),
        ("Söndag", "8 dec", "Ledig", "Ledig"),
    ]
    doc = create_weekly_schedule(49, "2 dec", "8 dec 2025", week49_days)
    doc.save(os.path.join(output_dir, "SCHEMA_Vecka49_2-8dec.docx"))
    print("✅ SCHEMA Vecka 49 skapad!")
    
    # Week 50: 9-14 dec (Mon-Sun, but trip ends 12 dec)
    week50_days = [
        ("Måndag", "9 dec", "3 tim", "3 tim"),
        ("Tisdag", "10 dec", "3 tim", "3 tim"),
        ("Onsdag", "11 dec", "3 tim", "3 tim"),
        ("Torsdag", "12 dec", "3 tim", "3 tim"),
        ("Fredag", "13 dec", "Ledig", "Ledig"),
        ("Lördag", "14 dec", "Ledig", "Ledig"),
    ]
    doc = create_weekly_schedule(50, "9 dec", "14 dec 2025", week50_days)
    doc.save(os.path.join(output_dir, "SCHEMA_Vecka50_9-14dec.docx"))
    print("✅ SCHEMA Vecka 50 skapad!")

# =============================================================================
# 5. INTYG - Employment Certificate
# =============================================================================
def create_intyg():
    doc = Document()
    set_margins(doc, 2)
    
    title = doc.add_heading('INTYG OM ANSTÄLLNINGSFÖRHÅLLANDE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Company
    p = doc.add_paragraph()
    p.add_run("ARBETSGIVARE: ").bold = True
    p.add_run("Sami's Jackets AB (559489-5301)")
    
    p = doc.add_paragraph()
    p.add_run("ADRESS: ").bold = True
    p.add_run("Kungsgatan 24, 632 18 Eskilstuna")
    
    doc.add_paragraph()
    
    # Statement
    doc.add_heading('INTYGAS HÄRMED', level=1)
    
    doc.add_paragraph("Följande personer är anställda vid Sami's Jackets AB som timanställda butiksbiträden:")
    
    doc.add_paragraph()
    
    # Employees table
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Namn"
    headers[1].text = "Personnummer"
    headers[2].text = "Anställd sedan"
    headers[3].text = "Timlön"
    
    cells = table.rows[1].cells
    cells[0].text = "Emelie Nilsson"
    cells[1].text = "19990408-5306"
    cells[2].text = "2024-09-01"
    cells[3].text = "110 SEK"
    
    cells = table.rows[2].cells
    cells[0].text = "Mari Johansson"
    cells[1].text = "19720522-1646"
    cells[2].text = "2024-09-01"
    cells[3].text = "110 SEK"
    
    doc.add_paragraph()
    
    # Purpose
    doc.add_heading('SYFTE', level=1)
    doc.add_paragraph("Detta intyg utfärdas på begäran av de anställda för att styrka anställningsförhållandet.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    doc.add_heading('UNDERSKRIFT', level=1)
    
    doc.add_paragraph()
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Mohammad Sami Alsharef")
    doc.add_paragraph("VD, Sami's Jackets AB")
    doc.add_paragraph()
    doc.add_paragraph("Datum: _______________     Ort: Eskilstuna")
    
    doc.save(os.path.join(output_dir, "INTYG_Anstallning_Emelie_Mari.docx"))
    print("✅ INTYG skapad!")

# Run all
if __name__ == "__main__":
    print("\n📄 Skapar Word-dokument...\n")
    create_fullmakt()
    create_vanskapsavtal()
    create_nodkontakter()
    create_all_weekly_schedules()
    create_intyg()
    print(f"\n✅ Alla dokument sparade i:\n{output_dir}")
    print("\n📋 Skapade filer:")
    for f in os.listdir(output_dir):
        print(f"   • {f}")
